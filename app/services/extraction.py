"""Text extraction from PDF and DOCX uploads."""

import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

import docx
import pypdf

from app.config import DOCX_CONTENT_TYPE, PDF_CONTENT_TYPE, Settings, get_settings
from app.core.exceptions import ExtractionError, UnsupportedFileTypeError
from app.core.logging import get_logger

#: Three or more consecutive newlines collapse to two, preserving paragraph breaks
#: while discarding runs of blank lines.
_EXCESS_BLANK_LINES = re.compile(r"\n{3,}")
#: Trailing spaces/tabs left behind by PDF layout extraction.
_TRAILING_SPACES = re.compile(r"[ \t]+$", re.MULTILINE)
#: Joins consecutive pages. A blank line, so it reads as a paragraph break to the
#: splitter rather than introducing a separator it would treat specially.
PAGE_SEPARATOR = "\n\n"

log = get_logger(__name__)


@dataclass(frozen=True)
class ExtractedDocument:
    """The text recovered from a document, plus what we learned while reading it."""

    text: str
    page_count: int | None
    char_count: int
    #: Character offset in ``text`` where each page begins, one entry per page. This is
    #: what lets a chunk be cited to the page it actually came from rather than to the
    #: document's total page count. ``None`` for formats without pages.
    page_offsets: list[int] | None = None


def _normalize(text: str) -> str:
    """Tidy extracted text without destroying its paragraph structure."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _TRAILING_SPACES.sub("", text)
    text = _EXCESS_BLANK_LINES.sub("\n\n", text)
    return text.strip()


def _fail(message: str, exc: Exception, **context: object) -> ExtractionError:
    """Log the underlying error and return a sanitized one for the client.

    Third-party parser messages routinely embed absolute filesystem paths, so they are
    recorded server-side only and never interpolated into the response body.
    """
    log.warning("extract.failed", error=type(exc).__name__, detail=str(exc), **context)
    return ExtractionError(message)


def _check_zip_bomb(file_path: Path, settings: Settings) -> None:
    """Reject DOCX archives that would expand to an unreasonable size.

    python-docx materializes each zip member in memory, so a highly compressible
    document well under the upload limit can otherwise exhaust the process.
    """
    try:
        with zipfile.ZipFile(file_path) as archive:
            uncompressed = sum(item.file_size for item in archive.infolist())
    except (zipfile.BadZipFile, OSError) as exc:
        raise _fail("The DOCX file could not be parsed.", exc, format="docx") from exc

    compressed = file_path.stat().st_size
    ratio_exceeded = compressed > 0 and uncompressed > compressed * settings.max_compression_ratio
    if uncompressed > settings.max_uncompressed_bytes or ratio_exceeded:
        log.warning(
            "extract.rejected",
            reason="decompression_bomb",
            compressed_bytes=compressed,
            uncompressed_bytes=uncompressed,
        )
        raise ExtractionError("The document failed structural validation.")


def _extract_pdf(file_path: Path) -> tuple[str, int, list[int]]:
    """Return the text of every page, the page count, and where each page starts.

    Pages are normalized individually and then joined, so the offsets stay valid: a
    single normalization pass over the joined text would shift every position and make
    the page mapping quietly wrong.
    """
    try:
        reader = pypdf.PdfReader(file_path)
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # pypdf raises a wide range of errors on malformed input
        raise _fail("The PDF file could not be parsed.", exc, format="pdf") from exc

    parts: list[str] = []
    offsets: list[int] = []
    cursor = 0
    for page_text in pages:
        normalized = _normalize(page_text)
        if parts and normalized:
            parts.append(PAGE_SEPARATOR)
            cursor += len(PAGE_SEPARATOR)
        # Recorded even for a blank page, so offsets stay one-per-page and monotonic.
        offsets.append(cursor)
        if normalized:
            parts.append(normalized)
            cursor += len(normalized)
    return "".join(parts), len(pages), offsets


def _extract_docx(file_path: Path, settings: Settings) -> str:
    """Return the document's paragraph text joined by newlines."""
    _check_zip_bomb(file_path, settings)
    try:
        document = docx.Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
    except Exception as exc:  # python-docx surfaces zip/XML errors on malformed input
        raise _fail("The DOCX file could not be parsed.", exc, format="docx") from exc
    return "\n".join(paragraphs)


def extract_text(
    file_path: Path, content_type: str, settings: Settings | None = None
) -> ExtractedDocument:
    """Extract normalized text from ``file_path``, dispatching on ``content_type``.

    This is synchronous and CPU-bound; callers on the event loop must run it in a
    thread pool.

    Raises:
        UnsupportedFileTypeError: no extractor is registered for the content type.
        ExtractionError: the file could not be parsed, or contained no text.
    """
    settings = settings or get_settings()

    page_offsets: list[int] | None = None
    if content_type == PDF_CONTENT_TYPE:
        # Already normalized per page, so offsets remain valid; do not re-normalize.
        text, page_count, page_offsets = _extract_pdf(file_path)
    elif content_type == DOCX_CONTENT_TYPE:
        text, page_count = _normalize(_extract_docx(file_path, settings)), None
    else:
        raise UnsupportedFileTypeError(f"Unsupported content type: {content_type or 'unknown'}")

    if not text:
        # A structurally valid file we cannot use is as much a failure as a corrupt one:
        # downstream retrieval has nothing to index.
        raise ExtractionError("No text could be extracted from the document.")

    return ExtractedDocument(
        text=text,
        page_count=page_count,
        char_count=len(text),
        page_offsets=page_offsets,
    )
